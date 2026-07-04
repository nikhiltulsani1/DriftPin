from __future__ import annotations

from pathlib import Path

import pypdf
import pytest
from docx import Document as DocxDocument

from driftpin.ingestion.parsers import UnsupportedDocumentFormatError, parse_document


def test_parse_markdown_splits_into_paragraph_blocks(tmp_path: Path) -> None:
    path = tmp_path / "prd.md"
    path.write_text(
        "The system shall allow password reset via email.\n\n"
        "Sessions must expire after 30 minutes of inactivity.\n",
        encoding="utf-8",
    )

    blocks = parse_document(path)

    assert len(blocks) == 2
    assert "password reset" in blocks[0].text
    assert blocks[0].anchor.startswith("lines ")
    assert blocks[0].source_doc_path == str(path)


def test_parse_txt_treats_blank_lines_as_block_separators(tmp_path: Path) -> None:
    path = tmp_path / "prd.txt"
    path.write_text("Requirement one text.\n\n\nRequirement two text.\n", encoding="utf-8")

    blocks = parse_document(path)

    assert len(blocks) == 2
    assert blocks[1].text == "Requirement two text."


def test_parse_docx_extracts_nonempty_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "prd.docx"
    document = DocxDocument()
    document.add_paragraph("Users must be able to reset passwords.")
    document.add_paragraph("")
    document.add_paragraph("Admins can lock accounts after 5 failed attempts.")
    document.save(str(path))

    blocks = parse_document(path)

    assert len(blocks) == 2
    assert blocks[0].anchor == "paragraph 1"
    assert "reset passwords" in blocks[0].text


def test_parse_pdf_runs_without_error_on_a_real_pdf(tmp_path: Path) -> None:
    path = tmp_path / "prd.pdf"
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with path.open("wb") as handle:
        writer.write(handle)

    blocks = parse_document(path)

    assert blocks == []  # blank page has no extractable text, but parsing must not raise


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    path = tmp_path / "prd.xyz"
    path.write_text("irrelevant", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentFormatError):
        parse_document(path)
